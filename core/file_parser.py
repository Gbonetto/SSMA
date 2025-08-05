import os
import io
import logging
from typing import Tuple, Dict, Optional

import pdfplumber
import docx
import docx2txt
import openpyxl
import xlrd
import pytesseract
from PIL import Image

try:
    import chardet  # Optionnel, mais conseillé !
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False

SUPPORTED_FORMATS = [
    "pdf", "docx", "doc", "txt", "xlsx", "xls", "png", "jpg", "jpeg"
]

# Configure logging (toujours UTF-8)
logging.basicConfig(level=logging.INFO, encoding='utf-8')

def get_extension(filename):
    return filename.split(".")[-1].lower()

def extract_text_and_metadata(filename: str, content: bytes) -> Tuple[Optional[str], Dict]:
    ext = get_extension(filename)
    text, meta = None, {}

    try:
        if ext == "pdf":
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages_text = [page.extract_text() or "" for page in pdf.pages]
                text = "\n".join(pages_text)
                meta = {"nb_pages": len(pdf.pages)}
            # Fallback OCR si PDF scanné
            if not text.strip():
                from pdf2image import convert_from_bytes
                pages = convert_from_bytes(content)
                text = "\n".join([pytesseract.image_to_string(p.convert("RGB")) for p in pages])
                meta["ocr"] = True

        elif ext == "docx":
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            meta["nb_paragraphs"] = len(doc.paragraphs)

        elif ext == "doc":
            with open("temp_doc.doc", "wb") as f:
                f.write(content)
            text = docx2txt.process("temp_doc.doc")
            os.remove("temp_doc.doc")
            meta["legacy_doc"] = True

        elif ext == "txt":
            # Encodage "intelligent"
            if HAS_CHARDET:
                detected = chardet.detect(content)
                enc = detected['encoding'] or 'utf-8'
                try:
                    text = content.decode(enc)
                    meta["encoding_used"] = enc
                except Exception:
                    text = content.decode("utf-8", errors="replace")
                    meta["encoding_fallback"] = "utf-8-replace"
            else:
                try:
                    text = content.decode("utf-8")
                    meta["encoding_used"] = "utf-8"
                except UnicodeDecodeError:
                    try:
                        text = content.decode("latin-1")
                        meta["encoding_used"] = "latin-1"
                    except Exception:
                        text = content.decode("utf-8", errors="replace")
                        meta["encoding_fallback"] = "utf-8-replace"

        elif ext == "xlsx":
            wb = openpyxl.load_workbook(io.BytesIO(content))
            text = ""
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_str = " ".join([str(cell) for cell in row if cell is not None])
                    if row_str.strip():
                        text += row_str + "\n"
            meta["nb_sheets"] = len(wb.worksheets)

        elif ext == "xls":
            wb = xlrd.open_workbook(file_contents=content)
            text = ""
            for sheet in wb.sheets():
                for row_idx in range(sheet.nrows):
                    row_str = " ".join([str(cell) for cell in sheet.row_values(row_idx) if cell is not None])
                    if row_str.strip():
                        text += row_str + "\n"
            meta["nb_sheets"] = len(wb.sheets())

        elif ext in ["jpg", "jpeg", "png"]:
            image = Image.open(io.BytesIO(content)).convert("RGB")
            text = pytesseract.image_to_string(image)
            meta["img_size"] = image.size

        else:
            logging.warning(f"Format non supporté : {filename}")
            return None, {"filename": filename, "ext": ext, "error": "Format non supporté"}

    except Exception as e:
        logging.error(f"Erreur parsing {filename} ({ext}): {e}")
        return None, {"filename": filename, "ext": ext, "error": str(e)}

    meta["filename"] = filename
    meta["ext"] = ext
    return text, meta
